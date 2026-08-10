"""Render confirmed chapter content into a copy of the Word template.

This walks each chapter's `content_json` (the same ProseMirror JSON that
Tiptap renders on screen in ContentPanel.vue) node by node, so the exported
DOCX reproduces the same headings, bold/highlight runs, bullet lists and
tables the user sees in the editor — instead of the previous behavior of
dumping `chapter.plain_text` (raw markdown-ish text with literal "##",
"**bold**", "| a | b |" that was never parsed) as flat unstyled paragraphs.
"""
import json
import re
import shutil
import unicodedata
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Twips
from app.config import get_storage_path
from app.services.export_fonts import get_export_font_config

EXPORT_FONTS = get_export_font_config()
DEFAULT_LATIN_FONT = EXPORT_FONTS.latin
DEFAULT_EAST_ASIA_FONT = EXPORT_FONTS.cjk
DEFAULT_COMPLEX_SCRIPT_FONT = EXPORT_FONTS.cjk


def render_to_docx(doc_id: str, chapters: list, template_source_path: str | None) -> str:
    """Write confirmed chapter content into a copy of the Word template."""
    out_dir = get_storage_path("generated") / doc_id
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "output.docx"

    if template_source_path and Path(template_source_path).exists():
        shutil.copy2(template_source_path, out_path)
        doc = Document(str(out_path))
        _inject_into_template(doc, chapters)
    else:
        doc = Document()
        _build_from_scratch(doc, chapters)

    doc.save(str(out_path))
    return str(out_path)


def _inject_into_template(doc: Document, chapters: list):
    """Find each chapter's heading paragraph in the template and insert the
    chapter's rendered content (headings/paragraphs/lists/tables) after it."""
    anchors = _template_chapter_anchors(doc, chapters)
    if anchors:
        _remove_template_placeholders(doc, anchors)
        for chapter, heading_para in anchors:
            _insert_nodes_after(doc, heading_para, _nodes_for_template(chapter))

    _ensure_template_footer_page_fields(doc)


def _build_from_scratch(doc: Document, chapters: list):
    """Fallback used when the target template file is unavailable."""
    _configure_generated_document(doc, chapters)
    for chapter in chapters:
        heading = doc.add_heading(chapter.title, level=1)
        _style_generated_heading(heading)
        for node in _chapter_nodes(chapter):
            _append_node(doc, node)


def _chapter_nodes(chapter) -> list[dict]:
    """Return the ProseMirror content nodes for a chapter, plus a trailing
    highlighted note for any missing-information / conflict items — mirroring
    exactly what ContentPanel.vue shows on screen (missingItems/conflictItems
    alert boxes + the Tiptap-rendered content_json body)."""
    nodes: list[dict] = []
    content_json = getattr(chapter, "content_json", None)
    if content_json:
        try:
            parsed = json.loads(content_json)
        except (TypeError, json.JSONDecodeError, AttributeError):
            pass
        else:
            content = parsed.get("content") if isinstance(parsed, dict) else None
            if isinstance(content, list):
                nodes.extend(
                    sanitized
                    for node in content
                    if (sanitized := _sanitize_block_node(node)) is not None
                )
    plain_text = getattr(chapter, "plain_text", None)
    if not nodes and isinstance(plain_text, str) and plain_text:
        nodes.append({"type": "paragraph", "content": [{"type": "text", "text": plain_text}]})

    conflicts = _safe_json_list(getattr(chapter, "conflict_json", None))
    if conflicts:
        descriptions = [_conflict_detail(conflict) for conflict in conflicts]
        descriptions = [description for description in descriptions if description]
        if descriptions:
            nodes.append(_notice_node("内容冲突：", '; '.join(descriptions[:5]), kind="conflict"))

    if not nodes:
        nodes.append({"type": "paragraph", "content": [{"type": "text", "text": "（内容待生成）"}]})
    return nodes


def _conflict_detail(entry) -> str:
    if isinstance(entry, dict):
        detail = (
            entry.get("description")
            or entry.get("title")
            or entry.get("name")
            or json.dumps(entry, ensure_ascii=False)
        )
    elif entry is None:
        return ""
    else:
        detail = str(entry).strip()
    return str(detail).strip()


def _notice_node(label: str, detail: str, kind: str = "missing") -> dict:
    return {
        "type": "paragraph",
        "attrs": {"notice_kind": kind},
        "content": [
            {"type": "text", "marks": [{"type": "highlight"}, {"type": "bold"}], "text": f"【{label}"},
            {"type": "text", "marks": [{"type": "highlight"}], "text": f"{detail}】"},
        ],
    }


def _node_content(node: dict) -> list | None:
    if not isinstance(node, dict):
        return None
    if "content" not in node:
        return []
    content = node.get("content")
    return content if isinstance(content, list) else None


def _node_attrs(node: dict) -> dict:
    if not isinstance(node, dict):
        return {}
    attrs = node.get("attrs")
    return attrs if isinstance(attrs, dict) else {}


def _sanitize_block_node(node) -> dict | None:
    if not isinstance(node, dict):
        return None

    node_type = node.get("type")
    if not isinstance(node_type, str):
        return None

    if node_type in ("paragraph", "heading"):
        content = _node_content(node)
        if content is None:
            return None
        runs = _sanitize_inline_content(content)
        if content and not runs:
            return None
        return {"type": node_type, "attrs": _node_attrs(node), "content": runs}

    if node_type in ("bulletList", "orderedList"):
        content = _node_content(node)
        if content is None:
            return None
        items = [_sanitize_list_item(item) for item in content]
        items = [item for item in items if item is not None]
        if not items:
            return None
        return {"type": node_type, "attrs": _node_attrs(node), "content": items}

    if node_type == "table":
        content = _node_content(node)
        if content is None:
            return None
        rows = [_sanitize_table_row(row) for row in content]
        rows = [row for row in rows if row is not None]
        if not rows:
            return None
        return {"type": node_type, "attrs": _node_attrs(node), "content": rows}

    # Keep unsupported but well-formed node types on the existing empty-
    # paragraph path. Their malformed children are intentionally discarded.
    return {"type": node_type}


def _sanitize_inline_content(content: list) -> list[dict]:
    runs = []
    for run_node in content:
        if not isinstance(run_node, dict) or run_node.get("type") != "text":
            continue
        text = run_node.get("text")
        if not isinstance(text, str) or not text:
            continue

        raw_marks = run_node.get("marks")
        marks = []
        if isinstance(raw_marks, list):
            marks = [
                {"type": mark_type}
                for mark in raw_marks
                if isinstance(mark, dict)
                and isinstance(mark_type := mark.get("type"), str)
            ]
        runs.append({"type": "text", "text": text, "marks": marks})
    return runs


def _sanitize_list_item(node) -> dict | None:
    if not isinstance(node, dict) or node.get("type") != "listItem":
        return None
    content = _node_content(node)
    if content is None:
        return None

    inner_nodes = []
    for inner in content:
        sanitized = _sanitize_block_node(inner)
        if sanitized and sanitized["type"] in ("paragraph", "bulletList", "orderedList"):
            inner_nodes.append(sanitized)
    if not inner_nodes:
        return None
    return {"type": "listItem", "attrs": _node_attrs(node), "content": inner_nodes}


def _sanitize_table_row(node) -> dict | None:
    if not isinstance(node, dict) or node.get("type") != "tableRow":
        return None
    content = _node_content(node)
    if content is None:
        return None

    cells = [_sanitize_table_cell(cell) for cell in content]
    cells = [cell for cell in cells if cell is not None]
    if not cells:
        return None
    return {"type": "tableRow", "attrs": _node_attrs(node), "content": cells}


def _sanitize_table_cell(node) -> dict | None:
    if not isinstance(node, dict) or node.get("type") not in ("tableCell", "tableHeader"):
        return None
    content = _node_content(node)
    if content is None:
        return None

    blocks = []
    for block in content:
        sanitized = _sanitize_block_node(block)
        if sanitized and sanitized["type"] in ("paragraph", "heading"):
            blocks.append(sanitized)
    if content and not blocks:
        return None
    return {"type": node["type"], "attrs": _node_attrs(node), "content": blocks}


def _insert_nodes_after(doc: Document, heading_para, nodes: list[dict]):
    """Render each node by appending it to the end of `doc` (via the normal
    high-level python-docx API, so styles/borders/runs are built correctly),
    then detach it from the body and re-insert it right after the heading in
    document order."""
    anchor = heading_para._element
    for node in nodes:
        for element in _render_node_elements(doc, node):
            element.getparent().remove(element)
            anchor.addnext(element)
            anchor = element


def _append_node(doc: Document, node: dict):
    for element in _render_node_elements(doc, node):
        pass  # already appended to doc; nothing further to do in from-scratch mode


def _render_node_elements(doc: Document, node: dict) -> list:
    """Append `node` to the end of `doc` using python-docx's high-level API
    and return the created XML element(s), in document order."""
    node = _sanitize_block_node(node)
    if node is None:
        return [doc.add_paragraph()._element]
    node_type = node.get("type")

    if node_type == "heading":
        level = _heading_node_level(node)
        para = doc.add_paragraph(style=f"Heading {level}")
        _add_runs(para, node.get("content", []))
        _style_generated_heading(para)
        return [para._element]

    if node_type == "paragraph":
        para = doc.add_paragraph()
        _add_runs(para, node.get("content", []))
        _style_paragraph(para)
        notice_kind = _node_attrs(node).get("notice_kind")
        if isinstance(notice_kind, str) and notice_kind:
            _style_notice(para, notice_kind)
        return [para._element]

    if node_type in ("bulletList", "orderedList"):
        return _render_list(doc, node, ordered=node_type == "orderedList")

    if node_type == "table":
        return [_render_table(doc, node)]

    # Unsupported node types (e.g. horizontalRule) fall back to an empty
    # paragraph rather than raising, so one odd node never breaks the export.
    return [doc.add_paragraph()._element]


def _render_list(doc: Document, node: dict, ordered: bool) -> list:
    node = _sanitize_block_node(node)
    if node is None:
        return []
    return _render_list_level(doc, node, ordered=ordered, level=0, number_prefix=())


def _render_list_level(
    doc: Document,
    node: dict,
    ordered: bool,
    level: int,
    number_prefix: tuple[int, ...],
) -> list:
    elements = []
    if not isinstance(node, dict):
        return elements
    attrs = _node_attrs(node)
    start = _list_start(attrs) if ordered else 1

    content = _node_content(node) or []
    for item_index, item in enumerate(content):
        if not isinstance(item, dict) or item.get("type") != "listItem":
            continue
        item_number = start + item_index
        paragraph_index = 0
        for inner in _node_content(item) or []:
            if not isinstance(inner, dict):
                continue
            inner_type = inner.get("type")
            if inner_type == "paragraph":
                para = doc.add_paragraph()
                if ordered and paragraph_index == 0:
                    full_number = (*number_prefix, item_number)
                    prefix = f"{'.'.join(str(number) for number in full_number)}. "
                elif ordered:
                    prefix = "   "
                elif paragraph_index == 0:
                    prefix = f"{_bullet_marker(level)} "
                else:
                    prefix = "  "
                list_run = para.add_run(prefix)
                _style_run_font(list_run)
                _add_runs(para, inner.get("content", []))
                _style_list_paragraph(para, level)
                elements.append(para._element)
                paragraph_index += 1
            elif inner_type in ("bulletList", "orderedList"):
                nested_prefix = (*number_prefix, item_number) if ordered else ()
                elements.extend(
                    _render_list_level(
                        doc,
                        inner,
                        ordered=inner_type == "orderedList",
                        level=level + 1,
                        number_prefix=nested_prefix,
                    )
                )
    return elements


def _heading_node_level(node: dict) -> int:
    raw_level = _node_attrs(node).get("level", 1)
    try:
        return max(1, min(int(raw_level), 3))
    except (TypeError, ValueError):
        return 1


def _bullet_marker(level: int) -> str:
    try:
        safe_level = max(0, int(level))
    except (TypeError, ValueError):
        safe_level = 0
    return ("•", "◦", "▪", "▫")[min(safe_level, 3)]


def _style_list_paragraph(paragraph, level: int):
    _style_paragraph(paragraph)
    paragraph.paragraph_format.left_indent = Mm(6 * level)
    paragraph.paragraph_format.first_line_indent = Mm(-4)


def _list_start(attrs: dict) -> int:
    if not isinstance(attrs, dict):
        return 1
    try:
        return int(attrs.get("start", attrs.get("order", 1)))
    except (TypeError, ValueError):
        return 1


def _add_runs(paragraph, inline_content: list[dict]):
    if not isinstance(inline_content, list):
        return
    for run_node in inline_content:
        if not isinstance(run_node, dict) or run_node.get("type") != "text":
            continue
        text = run_node.get("text", "")
        if not isinstance(text, str) or not text:
            continue
        raw_marks = run_node.get("marks", [])
        marks = {
            mark_type
            for mark in raw_marks
            if isinstance(mark, dict)
            and isinstance(mark_type := mark.get("type"), str)
        } if isinstance(raw_marks, list) else set()
        run = paragraph.add_run(text)
        _style_run_font(run)
        if "bold" in marks:
            run.bold = True
        if "italic" in marks:
            run.italic = True
        if "highlight" in marks:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _render_table(doc: Document, node: dict):
    node = _sanitize_block_node(node)
    rows = node.get("content", []) if node is not None else []
    if not rows:
        return doc.add_paragraph()._element

    n_cols = max((len(row.get("content", [])) for row in rows), default=0)
    if not n_cols:
        return doc.add_paragraph()._element
    table = doc.add_table(rows=len(rows), cols=n_cols)
    _set_table_borders(table)

    for r, row_node in enumerate(rows):
        row_cells = row_node.get("content", [])
        is_header_row = bool(row_cells) and row_cells[0].get("type") == "tableHeader"
        for c, cell_node in enumerate(row_cells):
            cell = table.cell(r, c)
            cell.text = ""
            cell_para = cell.paragraphs[0]
            _style_paragraph(cell_para)
            for cell_content in cell_node.get("content", []):
                _add_runs(cell_para, cell_content.get("content", []))
            if is_header_row:
                for run in cell_para.runs:
                    run.bold = True
                _shade_cell(cell, "D9E2F3")

    _style_table(table)
    return table._element


def _style_table(table):
    """Apply borders and bounded proportional widths.

    Widths are distributed from a fixed printable width budget so narrow
    identifier columns stay compact while longer text columns receive more
    room without depending on renderer-specific autofit heuristics.
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    total_width_twips = Cm(16.5).twips
    col_widths = _proportional_column_widths(table, total_width_twips)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Twips(col_widths[index])
            _set_cell_width(cell, col_widths[index])
            _set_cell_margins(cell, top=90, start=100, bottom=90, end=100)


def _set_table_borders(table):
    """The scenario1 target template has no 'Table Grid' style defined, so
    borders are applied directly via OOXML (0.5pt on every edge) to match
    the visual table styling used in ContentPanel.vue's word-body CSS."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "333333")
        borders.append(el)
    tbl_pr.append(borders)


def _nodes_for_template(chapter) -> list[dict]:
    nodes = _chapter_nodes(chapter)
    while nodes and nodes[0].get("type") == "heading":
        heading_text = _normalized_heading_text(_node_text(nodes[0]))
        chapter_title = _normalized_heading_text(getattr(chapter, "title", ""))
        if heading_text and heading_text == chapter_title:
            nodes = nodes[1:]
            continue
        break
    return nodes


def _template_chapter_anchors(doc: Document, chapters: list) -> list[tuple[object, object]]:
    paragraphs = list(doc.paragraphs)
    used_paragraphs = set()
    anchors = []

    for chapter in chapters:
        title = str(getattr(chapter, "title", "") or "").strip()
        if not title:
            continue
        normalized_title = _normalized_heading_text(title)
        for para in paragraphs:
            paragraph_key = id(para._element)
            if paragraph_key in used_paragraphs:
                continue
            if (
                normalized_title
                and _heading_text_matches(title, para.text)
                and para.style.name.startswith("Heading")
            ):
                used_paragraphs.add(paragraph_key)
                anchors.append((chapter, para))
                break

    body = doc.element.body
    return sorted(anchors, key=lambda item: list(body).index(item[1]._element))


def _remove_template_placeholders(doc: Document, anchors: list[tuple[object, object]]):
    """Remove replaceable template body while retaining document structure.

    Only body regions after matched chapter anchors are cleaned. Heading
    paragraphs, captions, tables, section boundaries, and drawing paragraphs
    remain structural unless a table/caption carries an explicit sample marker
    or is a proven duplicate of a generated chapter table. Cover content before
    the first anchor and header/footer parts are never touched here.
    """
    paragraphs_by_element = {para._element: para for para in doc.paragraphs}
    for index, (chapter, heading_para) in enumerate(anchors):
        next_heading = (
            anchors[index + 1][1]._element
            if index + 1 < len(anchors)
            else None
        )
        replacement_elements = _template_table_replacement_elements(
            chapter,
            heading_para,
            next_heading,
            paragraphs_by_element,
        )
        remove_next_table = False
        current = heading_para._element.getnext()
        while current is not None and current != next_heading:
            next_element = current.getnext()
            if not _is_template_structure_element(current):
                paragraph = paragraphs_by_element.get(current)
                if paragraph is not None:
                    if current in replacement_elements:
                        current.getparent().remove(current)
                    elif _paragraph_has_drawing(paragraph):
                        pass
                    elif _is_heading_paragraph(paragraph):
                        if _is_sample_heading(paragraph.text):
                            current.getparent().remove(current)
                    elif _is_caption_paragraph(paragraph):
                        if _is_sample_placeholder_text(paragraph.text):
                            remove_next_table = (
                                next_element is not None
                                and next_element.tag == qn("w:tbl")
                            )
                            current.getparent().remove(current)
                    elif _is_sample_placeholder_text(paragraph.text):
                        current.getparent().remove(current)
                elif current.tag == qn("w:tbl") and (
                    current in replacement_elements
                    or _should_remove_template_table(current, current.getprevious())
                    or remove_next_table
                ):
                    _remove_template_table_element(current)
                    remove_next_table = False
                elif current.tag == qn("w:tbl"):
                    remove_next_table = False
            current = next_element


def _is_template_structure_element(element) -> bool:
    if element.tag == qn("w:sectPr"):
        return True
    return element.find(".//" + qn("w:sectPr")) is not None


def _normalized_heading_text(value) -> str:
    normalized = unicodedata.normalize("NFKC", str(value or ""))
    return "".join(
        character
        for character in normalized
        if not character.isspace() and not unicodedata.category(character).startswith("P")
    )


def _heading_text_matches(title: str, paragraph_text: str) -> bool:
    normalized_title = _normalized_heading_text(title)
    normalized_paragraph = _normalized_heading_text(paragraph_text)
    if not normalized_title or not normalized_paragraph:
        return False
    if normalized_title == normalized_paragraph:
        return True
    if not normalized_paragraph.startswith(normalized_title):
        return False

    prefix_end = _normalized_prefix_end(paragraph_text, normalized_title)
    if prefix_end is None:
        return False
    suffix = str(paragraph_text or "")[prefix_end:]
    if not suffix:
        return True
    if suffix[0].isspace():
        return True
    first_suffix_character = next((character for character in suffix if not character.isspace()), "")
    return bool(
        first_suffix_character
        and unicodedata.category(first_suffix_character).startswith("P")
    )


def _normalized_prefix_end(value: str, normalized_prefix: str) -> int | None:
    normalized_so_far = ""
    for index, character in enumerate(str(value or "")):
        normalized_character = _normalized_heading_text(character)
        if not normalized_character:
            continue
        normalized_so_far += normalized_character
        if not normalized_prefix.startswith(normalized_so_far):
            return None
        if normalized_so_far == normalized_prefix:
            return index + 1
    return None


def _is_heading_paragraph(paragraph) -> bool:
    return paragraph.style.name.startswith("Heading")


def _is_sample_heading(text: str) -> bool:
    return bool(re.match(r"^\s*[一二三四五六七八九十]+级标题[（(]", text or ""))


def _is_caption_paragraph(paragraph) -> bool:
    return paragraph.style.name == "Caption" or bool(
        re.match(r"^\s*[表图]\s*\d", paragraph.text or "")
    )


def _is_sample_placeholder_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact:
        return False
    if any(marker in compact for marker in ("占位", "示例", "样例", "XXXX", "例：", "例:")):
        return True
    if "**规格的**" in compact or "**（用途和适用范围）" in compact:
        return True
    if any(
        marker in compact
        for marker in (
            "填写说明",
            "来自附件",
            "待用户填写",
            "本章节要明确",
            "必要性总结",
            "满足需求、打压对手、呈现价值",
            "此表是贯穿",
            "文件内容填写要求",
            "附件中出现",
            "表题：",
            "表格框线：",
            "表格中内容：",
            "图题：",
            "图片居中",
            "图片与上方文字",
            "封面页的",
            "运输形式：",
            "运输方式：",
            "包装方式：",
        )
    ):
        return True
    return "模板" in compact and any(marker in compact for marker in ("说明", "末尾", "占位"))


def _paragraph_has_drawing(paragraph) -> bool:
    return paragraph._element.find(".//" + qn("w:drawing")) is not None


def _remove_template_table_element(element):
    """Remove a stale table without dropping drawings embedded in its cells."""
    drawing_paragraphs = [
        paragraph
        for paragraph in element.iter(qn("w:p"))
        if paragraph.find(".//" + qn("w:drawing")) is not None
    ]
    insertion_anchor = element
    for paragraph in drawing_paragraphs:
        paragraph.getparent().remove(paragraph)
        insertion_anchor.addprevious(paragraph)
        insertion_anchor = paragraph
    element.getparent().remove(element)


def _template_table_replacement_elements(
    chapter,
    heading_para,
    next_heading,
    paragraphs_by_element: dict[object, object],
) -> set[object]:
    """Select bounded table slots that generated tables replace.

    A generated table can replace a captioned table only inside the current
    heading's structural region, and at most one template table per generated
    table. Explicit placeholder/sample markers remain independently removable.
    Uncaptioned tables and tables in other heading regions are not candidates.
    """
    generated_table_count = sum(
        node.get("type") == "table" for node in _chapter_nodes(chapter)
    )
    region_end = _template_table_region_end(
        heading_para,
        next_heading,
        paragraphs_by_element,
    )
    candidates = []
    explicit_tables = set()
    current = heading_para._element.getnext()
    while current is not None and current != region_end:
        if current.tag == qn("w:tbl"):
            preceding = current.getprevious()
            preceding_para = paragraphs_by_element.get(preceding)
            explicit = _is_sample_placeholder_text(_element_text(current)) or (
                preceding is not None
                and _is_sample_placeholder_text(_element_text(preceding))
            )
            if explicit:
                explicit_tables.add(current)
            if preceding_para is not None and _is_caption_paragraph(preceding_para):
                candidates.append((current, preceding_para, explicit))
        current = current.getnext()

    replacement_elements = {
        table
        for table, _, _ in candidates
        if table in explicit_tables
    }
    for table, caption, explicit in candidates:
        if explicit:
            replacement_elements.add(caption._element)

    remaining = max(generated_table_count - len(explicit_tables), 0)
    if remaining == 0:
        return replacement_elements

    generated_captions = _generated_table_captions(chapter)
    available = [candidate for candidate in candidates if not candidate[2]]
    selected = []
    for generated_caption in generated_captions:
        if not generated_caption:
            continue
        match = next(
            (
                candidate
                for candidate in available
                if _captions_match(generated_caption, _element_text(candidate[1]._element))
            ),
            None,
        )
        if match is None:
            continue
        selected.append(match)
        available.remove(match)
        if len(selected) == remaining:
            break

    selected.extend(available[: max(remaining - len(selected), 0)])
    for table, caption, _ in selected:
        replacement_elements.add(table)
        replacement_elements.add(caption._element)
    return replacement_elements


def _template_table_region_end(
    heading_para,
    next_heading,
    paragraphs_by_element: dict[object, object],
):
    heading_level = _heading_level(heading_para)
    current = heading_para._element.getnext()
    while current is not None and current != next_heading:
        paragraph = paragraphs_by_element.get(current)
        if paragraph is not None and _is_heading_paragraph(paragraph):
            if _heading_level(paragraph) <= heading_level:
                return current
        if current.tag == qn("w:sectPr"):
            return current
        current = current.getnext()
    return next_heading


def _heading_level(paragraph) -> int:
    match = re.match(r"^Heading\s+(\d+)$", paragraph.style.name or "")
    return int(match.group(1)) if match else 99


def _generated_table_captions(chapter) -> list[str]:
    nodes = _chapter_nodes(chapter)
    captions = []
    for index, node in enumerate(nodes):
        if node.get("type") != "table":
            continue
        preceding = nodes[index - 1] if index else None
        captions.append(
            _node_text(preceding)
            if preceding is not None and preceding.get("type") == "paragraph"
            else ""
        )
    return captions


def _captions_match(left: str, right: str) -> bool:
    normalized_left = _normalized_heading_text(left)
    normalized_right = _normalized_heading_text(right)
    return bool(
        normalized_left
        and normalized_right
        and (
            normalized_left == normalized_right
            or normalized_left in normalized_right
            or normalized_right in normalized_left
        )
    )


def _should_remove_template_table(element, preceding_element) -> bool:
    if _is_sample_placeholder_text(_element_text(element)):
        return True
    if preceding_element is not None and preceding_element.tag == qn("w:p"):
        if _is_sample_placeholder_text(_element_text(preceding_element)):
            return True
    return False


def _element_text(element) -> str:
    if element is None:
        return ""
    return "".join(text.text or "" for text in element.iter(qn("w:t")))


def _configure_generated_document(doc: Document, chapters: list):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2))
    section.footer_distance = Cm(1.2)

    title = " / ".join(ch.title for ch in chapters if getattr(ch, "title", None))
    doc.core_properties.title = title or "导出文档"

    _configure_styles(doc)
    _ensure_footer_page_field(section)


def _configure_styles(doc: Document):
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = DEFAULT_LATIN_FONT
        r_pr = _ensure_child(style._element, "w:rPr")
        r_fonts = _ensure_child(r_pr, "w:rFonts")
        r_fonts.set(qn("w:ascii"), DEFAULT_LATIN_FONT)
        r_fonts.set(qn("w:hAnsi"), DEFAULT_LATIN_FONT)
        r_fonts.set(qn("w:eastAsia"), DEFAULT_EAST_ASIA_FONT)
        r_fonts.set(qn("w:cs"), DEFAULT_COMPLEX_SCRIPT_FONT)


def _style_run_font(run):
    run.font.name = DEFAULT_LATIN_FONT
    r_pr = _ensure_child(run._r, "w:rPr")
    r_fonts = _ensure_child(r_pr, "w:rFonts")
    r_fonts.set(qn("w:ascii"), DEFAULT_LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), DEFAULT_LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), DEFAULT_EAST_ASIA_FONT)
    r_fonts.set(qn("w:cs"), DEFAULT_COMPLEX_SCRIPT_FONT)


def _style_generated_heading(paragraph):
    paragraph.paragraph_format.space_before = Mm(6)
    paragraph.paragraph_format.space_after = Mm(3)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        _style_run_font(run)


def _style_paragraph(paragraph):
    paragraph.paragraph_format.space_after = Mm(2)


def _style_notice(paragraph, kind: str):
    paragraph.paragraph_format.space_before = Mm(2)
    paragraph.paragraph_format.space_after = Mm(2)
    p_pr = _ensure_child(paragraph._element, "w:pPr")
    _append_border_block(
        p_pr,
        "w:pBdr",
        color="C00000" if kind == "conflict" else "C9A227",
    )
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FDE9D9" if kind == "conflict" else "FFF2CC")


def _shade_cell(cell, fill: str):
    tc_pr = _ensure_child(cell._tc, "w:tcPr")
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int, start: int, bottom: int, end: int):
    tc_pr = _ensure_child(cell._tc, "w:tcPr")
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width: int):
    tc_pr = _ensure_child(cell._tc, "w:tcPr")
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _proportional_column_widths(table, total_width_twips: int):
    n_cols = max(len(row.cells) for row in table.rows) if table.rows else 1
    total_width_twips = max(int(total_width_twips), n_cols)
    text_weights = [1] * n_cols
    for col_index in range(n_cols):
        col_text_length = max(
            (len((row.cells[col_index].text or "").strip()) for row in table.rows if col_index < len(row.cells)),
            default=1,
        )
        text_weights[col_index] = max(1, min(col_text_length, 8))

    total_weight = sum(text_weights) or n_cols
    min_width = max(1, min(int(total_width_twips * 0.18), total_width_twips // n_cols))
    remaining_width = total_width_twips - (min_width * n_cols)
    widths = []
    for weight in text_weights:
        proportional_extra = int(remaining_width * weight / total_weight) if remaining_width else 0
        widths.append(min_width + proportional_extra)

    width_delta = total_width_twips - sum(widths)
    if widths and width_delta:
        widths[-1] += width_delta
    return widths


def _safe_json_list(raw_value) -> list:
    if not raw_value:
        return []
    try:
        parsed = json.loads(raw_value)
    except (TypeError, json.JSONDecodeError):
        return []
    return parsed if isinstance(parsed, list) else []


def _ensure_footer_page_field(section):
    footer = section.footer
    if _footer_has_page_field(footer):
        return

    paragraph = next(
        (candidate for candidate in footer.paragraphs if not _paragraph_has_content(candidate)),
        None,
    )
    if paragraph is None:
        style_name = footer.paragraphs[0].style.name if footer.paragraphs else None
        paragraph = footer.add_paragraph(style=style_name)
    paragraph.alignment = 1

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    _style_run_font(run)
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _footer_has_page_field(footer) -> bool:
    instructions = [
        (element.text or "")
        for element in footer._element.iter(qn("w:instrText"))
    ]
    instructions.extend(
        element.get(qn("w:instr"), "")
        for element in footer._element.iter(qn("w:fldSimple"))
    )
    return any(re.search(r"\bPAGE\b", instruction, re.IGNORECASE) for instruction in instructions)


def _ensure_template_footer_page_fields(doc: Document):
    for section in doc.sections:
        _ensure_footer_page_field(section)


def _paragraph_has_content(paragraph) -> bool:
    return any(child.tag != qn("w:pPr") for child in paragraph._element)


def _append_border_block(parent, tag_name: str, color: str):
    border_block = parent.find(qn(tag_name))
    if border_block is None:
        border_block = OxmlElement(tag_name)
        parent.append(border_block)
    for edge in ("top", "left", "bottom", "right"):
        edge_el = border_block.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            border_block.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "6")
        edge_el.set(qn("w:space"), "2")
        edge_el.set(qn("w:color"), color)


def _ensure_child(parent, tag_name: str):
    child = parent.find(qn(tag_name))
    if child is None:
        child = OxmlElement(tag_name)
        parent.append(child)
    return child


def _node_text(node: dict) -> str:
    if not isinstance(node, dict):
        return ""
    parts: list[str] = []
    for child in _node_content(node) or []:
        if not isinstance(child, dict):
            continue
        if child.get("type") == "text":
            text = child.get("text", "")
            if isinstance(text, str):
                parts.append(text)
        else:
            parts.append(_node_text(child))
    return "".join(parts)
