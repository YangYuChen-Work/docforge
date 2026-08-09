import platform
import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services import pdf_exporter
from app.services.export_fonts import get_export_font_config
from app.services.pdf_exporter import convert_to_pdf


def test_convert_to_pdf_raises_when_libreoffice_missing(monkeypatch, tmp_path):
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(b"fake-docx")

    monkeypatch.setattr("app.services.pdf_exporter._find_libreoffice", lambda: None)
    monkeypatch.setattr("app.services.pdf_exporter.settings.libreoffice_path", None)

    with pytest.raises(RuntimeError, match="LibreOffice 未找到"):
        convert_to_pdf(str(docx_path))


def test_convert_to_pdf_raises_clear_error_for_missing_configured_libreoffice(
    monkeypatch, tmp_path
):
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(b"fake-docx")
    monkeypatch.setattr(
        pdf_exporter.settings,
        "libreoffice_path",
        str(tmp_path / "missing" / "soffice"),
    )
    monkeypatch.setattr(
        pdf_exporter.subprocess,
        "run",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            FileNotFoundError(2, "No such file or directory")
        ),
    )

    with pytest.raises(RuntimeError, match="LibreOffice 未找到"):
        convert_to_pdf(str(docx_path))


def test_convert_to_pdf_gives_bundled_soffice_a_temporary_cjk_font_environment(
    monkeypatch, tmp_path
):
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(b"fake-docx")
    font_path = tmp_path / "PingFang.ttc"
    font_path.write_bytes(b"installed-font")
    base_config = tmp_path / "fonts.conf"
    base_config.write_text("<fontconfig><dir>/bundled/fonts</dir></fontconfig>")

    monkeypatch.setattr(pdf_exporter.settings, "libreoffice_path", "/bundled/soffice")
    monkeypatch.setattr(pdf_exporter.platform, "system", lambda: "Darwin")
    monkeypatch.setattr(pdf_exporter, "_find_cjk_font_file", lambda _font_name: font_path)
    monkeypatch.setattr(pdf_exporter, "_find_fontconfig_file", lambda _lo_bin: base_config)

    calls = []
    observed = {}

    def fake_run(command, **kwargs):
        calls.append((command, kwargs))
        font_config = Path(kwargs["env"]["FONTCONFIG_FILE"])
        observed["font_config"] = font_config.read_text()
        observed["font_bytes"] = (font_config.parent / "fonts" / font_path.name).read_bytes()
        docx_path.with_suffix(".pdf").write_bytes(b"fake-pdf")
        return subprocess.CompletedProcess(command, 0, stdout="", stderr="")

    monkeypatch.setattr(pdf_exporter.subprocess, "run", fake_run)

    assert convert_to_pdf(str(docx_path)) == str(docx_path.with_suffix(".pdf"))
    command, kwargs = calls[0]
    env = kwargs["env"]
    font_config = Path(env["FONTCONFIG_FILE"])

    assert command[0] == "/bundled/soffice"
    assert any(argument.startswith("-env:UserInstallation=file://") for argument in command)
    assert env["FONTCONFIG_PATH"] == str(base_config.parent)
    assert f"<dir>{font_config.parent / 'fonts'}</dir>" in observed["font_config"]
    assert observed["font_bytes"] == b"installed-font"
    assert not font_config.exists()


def test_fontconfig_fallback_retains_macos_and_bundled_font_directories(tmp_path):
    lo_contents = tmp_path / "LibreOffice.app" / "Contents"
    lo_bin = lo_contents / "MacOS" / "soffice"
    bundled_fonts = lo_contents / "Resources" / "fonts"
    bundled_fonts.mkdir(parents=True)
    font_directory = tmp_path / "temporary-fonts"

    config = pdf_exporter._fontconfig_with_font_directory(
        None, font_directory, str(lo_bin)
    )

    expected_directories = (
        font_directory,
        Path.home() / "Library" / "Fonts",
        Path("/Library/Fonts"),
        Path("/Network/Library/Fonts"),
        Path("/System/Library/Fonts"),
        Path("/System/Library/Fonts/Supplemental"),
        Path("/System/Library/AssetsV2"),
        Path("/System/Library/Assets"),
        bundled_fonts,
    )
    for directory in expected_directories:
        assert f"<dir>{directory}</dir>" in config


def test_temporary_cjk_environment_is_cleaned_after_subprocess_timeout(
    monkeypatch, tmp_path
):
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(b"fake-docx")
    font_path = tmp_path / "PingFang.ttc"
    font_path.write_bytes(b"installed-font")
    monkeypatch.setattr(pdf_exporter.settings, "libreoffice_path", "/bundled/soffice")
    monkeypatch.setattr(pdf_exporter.platform, "system", lambda: "Darwin")
    monkeypatch.setattr(pdf_exporter, "_find_cjk_font_file", lambda _font_name: font_path)
    monkeypatch.setattr(pdf_exporter, "_find_fontconfig_file", lambda _lo_bin: None)
    observed = {}

    def fake_run(command, **kwargs):
        observed["font_config"] = Path(kwargs["env"]["FONTCONFIG_FILE"])
        raise subprocess.TimeoutExpired(command, 60)

    monkeypatch.setattr(pdf_exporter.subprocess, "run", fake_run)

    with pytest.raises(subprocess.TimeoutExpired):
        convert_to_pdf(str(docx_path))

    assert not observed["font_config"].parent.exists()


def test_cleanup_failure_does_not_mask_nonzero_conversion_error(monkeypatch, tmp_path):
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(b"fake-docx")
    font_path = tmp_path / "PingFang.ttc"
    font_path.write_bytes(b"installed-font")
    monkeypatch.setattr(pdf_exporter.settings, "libreoffice_path", "/bundled/soffice")
    monkeypatch.setattr(pdf_exporter.platform, "system", lambda: "Darwin")
    monkeypatch.setattr(pdf_exporter, "_find_cjk_font_file", lambda _font_name: font_path)
    monkeypatch.setattr(pdf_exporter, "_find_fontconfig_file", lambda _lo_bin: None)
    cleanup_paths = []
    real_rmtree = shutil.rmtree

    def failing_rmtree(path):
        cleanup_paths.append(path)
        raise OSError("busy temporary directory")

    def fake_run(command, **kwargs):
        return subprocess.CompletedProcess(
            command, 1, stdout="", stderr="conversion failed"
        )

    monkeypatch.setattr(pdf_exporter.shutil, "rmtree", failing_rmtree)
    monkeypatch.setattr(pdf_exporter.subprocess, "run", fake_run)

    try:
        with pytest.raises(RuntimeError, match="LibreOffice 转换失败") as error:
            convert_to_pdf(str(docx_path))
        assert any("temporary PDF export cleanup failed" in note for note in error.value.__notes__)
    finally:
        monkeypatch.undo()
        for path in cleanup_paths:
            real_rmtree(path, ignore_errors=True)


def test_temporary_environment_reports_cleanup_failure_without_conversion_error(
    monkeypatch, tmp_path
):
    font_path = tmp_path / "PingFang.ttc"
    font_path.write_bytes(b"installed-font")
    monkeypatch.setattr(pdf_exporter.platform, "system", lambda: "Darwin")
    monkeypatch.setattr(pdf_exporter, "_find_cjk_font_file", lambda _font_name: font_path)
    cleanup_paths = []
    real_rmtree = shutil.rmtree

    def failing_rmtree(path):
        cleanup_paths.append(path)
        raise OSError("busy temporary directory")

    monkeypatch.setattr(pdf_exporter.shutil, "rmtree", failing_rmtree)

    try:
        with pytest.raises(RuntimeError, match="temporary PDF export cleanup failed"):
            with pdf_exporter._conversion_environment("/bundled/soffice"):
                pass
    finally:
        monkeypatch.undo()
        for path in cleanup_paths:
            real_rmtree(path, ignore_errors=True)


def test_convert_to_pdf_preserves_native_environment_off_macos(monkeypatch, tmp_path):
    docx_path = tmp_path / "sample.docx"
    docx_path.write_bytes(b"fake-docx")
    monkeypatch.setattr(pdf_exporter.settings, "libreoffice_path", "/system/soffice")
    monkeypatch.setattr(pdf_exporter.platform, "system", lambda: "Linux")

    calls = []

    def fake_run(command, **kwargs):
        calls.append((command, kwargs))
        docx_path.with_suffix(".pdf").write_bytes(b"fake-pdf")
        return subprocess.CompletedProcess(command, 0, stdout="", stderr="")

    monkeypatch.setattr(pdf_exporter.subprocess, "run", fake_run)

    convert_to_pdf(str(docx_path))

    command, kwargs = calls[0]
    assert not any(argument.startswith("-env:UserInstallation=") for argument in command)
    assert "env" not in kwargs


def _extract_pdf_text_candidates(pdf_path):
    candidates = []
    extractor_available = False
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            PdfReader = None
    if PdfReader is not None:
        extractor_available = True
        try:
            reader = PdfReader(str(pdf_path))
            candidates.append("\n".join(page.extract_text() or "" for page in reader.pages))
        except Exception:
            pass

    pdftotext_bin = shutil.which("pdftotext")
    if pdftotext_bin:
        extractor_available = True
        result = subprocess.run(
            [pdftotext_bin, str(pdf_path), "-"],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode == 0:
            candidates.append(result.stdout)
    return candidates, extractor_available


def _rendered_pdf_has_ink(pdf_path, tmp_path):
    pdftoppm_bin = shutil.which("pdftoppm")
    if not pdftoppm_bin:
        return None
    try:
        from PIL import Image, ImageChops
    except ImportError:
        return None

    output_prefix = tmp_path / "cjk-rendered"
    result = subprocess.run(
        [
            pdftoppm_bin,
            "-f",
            "1",
            "-l",
            "1",
            "-png",
            "-singlefile",
            str(pdf_path),
            str(output_prefix),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        return False

    try:
        with Image.open(output_prefix.with_suffix(".png")) as image:
            grayscale = image.convert("L")
            difference = ImageChops.difference(
                grayscale, Image.new("L", grayscale.size, color=255)
            )
            return difference.getbbox() is not None
    except OSError:
        return False


def _assert_pdf_smoke_result(
    expected_text, text_candidates, rendered_ink, extractor_available
):
    normalized_expected = "".join(expected_text.split())
    contains_expected_text = any(
        normalized_expected in "".join(candidate.split())
        for candidate in text_candidates
    )
    if extractor_available:
        assert contains_expected_text
    else:
        assert rendered_ink is True


def test_pdf_smoke_does_not_accept_rendered_ink_for_latin_only_output():
    with pytest.raises(AssertionError):
        _assert_pdf_smoke_result(
            expected_text="中文转换验证",
            text_candidates=["Latin-only output"],
            rendered_ink=True,
            extractor_available=True,
        )


@pytest.mark.skipif(platform.system() != "Darwin", reason="bundled CJK smoke test is macOS-specific")
def test_bundled_soffice_real_conversion_contains_or_renders_chinese_text(tmp_path):
    lo_bin = pdf_exporter.settings.libreoffice_path or pdf_exporter._find_libreoffice()
    if not lo_bin:
        pytest.skip("LibreOffice is required for the real PDF smoke test")

    docx_path = tmp_path / "cjk.docx"
    doc = Document()
    expected_text = "中文转换验证：产品概述、设计输入、质量目标。"
    run = doc.add_paragraph().add_run(expected_text)
    fonts = get_export_font_config()
    r_fonts = OxmlElement("w:rFonts")
    for attribute, value in (
        ("ascii", fonts.latin),
        ("hAnsi", fonts.latin),
        ("eastAsia", fonts.cjk),
        ("cs", fonts.cjk),
    ):
        r_fonts.set(qn(f"w:{attribute}"), value)
    run._r.get_or_add_rPr().append(r_fonts)
    doc.save(docx_path)

    pdf_path = Path(convert_to_pdf(str(docx_path)))
    text_candidates, extractor_available = _extract_pdf_text_candidates(pdf_path)
    rendered_ink = _rendered_pdf_has_ink(pdf_path, tmp_path)
    if not extractor_available and rendered_ink is None:
        pytest.skip("pypdf/PyPDF2, pdftotext, or pdftoppm with Pillow is required")

    _assert_pdf_smoke_result(
        expected_text,
        text_candidates,
        rendered_ink,
        extractor_available,
    )
