import json
from unittest.mock import MagicMock
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm
from app.services.docx_renderer import render_to_docx
from app.services.export_fonts import get_export_font_config
from app.services.validator import validate_document


def _ch(chapter_id, title, content_json, missing=None, conflicts=None):
    ch = MagicMock()
    ch.id = chapter_id
    ch.title = title
    ch.plain_text = "(unused fallback)"
    ch.content_json = json.dumps(content_json, ensure_ascii=False)
    ch.missing_information_json = json.dumps(missing or [], ensure_ascii=False)
    ch.conflict_json = json.dumps(conflicts or [], ensure_ascii=False)
    return ch


def _make_template(tmp_path, headings):
    """Build a minimal .docx with one Heading-1 paragraph per chapter title,
    mirroring the real target template's structure (heading, then body)."""
    doc = Document()
    for h in headings:
        doc.add_heading(h, level=1)
    path = tmp_path / "template.docx"
    doc.save(str(path))
    return str(path)


def _make_placeholder_template(tmp_path):
    doc = Document()
    doc.add_paragraph("模板封面标题")
    first_footer = doc.sections[0].footer
    first_footer.paragraphs[0].text = "模板首页页脚"
    first_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_section(WD_SECTION.NEW_PAGE)
    second_footer = doc.sections[1].footer
    second_footer.is_linked_to_previous = False
    second_footer.paragraphs[0].text = "模板正文页脚"
    second_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for title in ("第一章", "第二章"):
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"模板占位内容：{title}")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = f"模板占位表格：{title}"
    doc.add_paragraph("模板末尾说明")

    path = tmp_path / "placeholder-template.docx"
    doc.save(str(path))
    return str(path)


def test_heading_node_becomes_docx_heading_style(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "市场需求分析"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert any(p.text == "市场需求分析" for p in headings)


def test_bold_mark_becomes_bold_run(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "marks": [{"type": "bold"}], "text": "表6 产品卖点和优势分析表"},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
    assert any(r.text == "表6 产品卖点和优势分析表" for r in bold_runs)


def test_table_node_becomes_real_docx_table(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "table", "content": [
            {"type": "tableRow", "content": [
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "序号"}]}]},
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "类别"}]}]},
            ]},
            {"type": "tableRow", "content": [
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "1"}]}]},
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "臂架系统"}]}]},
            ]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "序号"
    assert table.rows[1].cells[1].text == "臂架系统"


def test_bullet_list_becomes_list_bullet_paragraphs(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "bulletList", "content": [
            {"type": "listItem", "content": [{"type": "paragraph", "content": [
                {"type": "text", "marks": [{"type": "bold"}], "text": "卖点"},
                {"type": "text", "text": "：主臂采用五节U型截面"},
            ]}]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    bullet_paras = [p for p in doc.paragraphs if "卖点" in p.text]
    assert len(bullet_paras) == 1
    assert "主臂采用五节U型截面" in bullet_paras[0].text


def test_ordered_list_becomes_numbered_paragraphs(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "orderedList", "content": [
            {"type": "listItem", "content": [{"type": "paragraph", "content": [
                {"type": "text", "text": "第一项"},
            ]}]},
            {"type": "listItem", "content": [{"type": "paragraph", "content": [
                {"type": "text", "text": "第二项"},
            ]}]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)

    numbered_paras = [p.text for p in doc.paragraphs if "项" in p.text]
    assert numbered_paras == ["1. 第一项", "2. 第二项"]


def test_template_cleanup_preserves_structure_and_footer_page_fields(tmp_path):
    content = lambda text: {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": text}]},
    ]}
    chapters = [
        _ch("c1", "第一章", content("生成第一章")),
        _ch("c2", "第二章", content("生成第二章")),
    ]
    template = _make_placeholder_template(tmp_path)

    out = render_to_docx("doc1", chapters, template)
    doc = Document(out)
    body_text = "\n".join(p.text for p in doc.paragraphs)

    assert "模板封面标题" in body_text
    assert "生成第一章" in body_text
    assert "生成第二章" in body_text
    assert "模板占位内容" not in body_text
    assert "模板末尾说明" not in body_text
    assert all(
        any(p.text == title and p.style.name == "Heading 1" for p in doc.paragraphs)
        for title in ("第一章", "第二章")
    )
    assert len(doc.tables) == 0

    first_footer = doc.sections[0].footer
    second_footer = doc.sections[1].footer
    first_footer_text = "\n".join(p.text for p in first_footer.paragraphs)
    second_footer_text = "\n".join(p.text for p in second_footer.paragraphs)
    assert "模板首页页脚" in first_footer_text
    assert "模板正文页脚" in second_footer_text
    assert first_footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert second_footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert "PAGE" in first_footer._element.xml
    assert "PAGE" in second_footer._element.xml


def test_missing_information_appended_as_highlighted_paragraph(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx(
        "doc1", [_ch("c1", "产品概述", content, missing=["下一年度销量预测"])], template
    )
    doc = Document(out)
    assert any("待补充" in p.text and "下一年度销量预测" in p.text for p in doc.paragraphs)


def test_missing_and_conflict_notes_are_rendered_as_highlighted_notices(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx(
        "doc1",
        [_ch(
            "c1",
            "产品概述",
            content,
            missing=["下一年度销量预测"],
            conflicts=[{"description": "销售目标与产能规划不一致"}],
        )],
        template,
    )

    doc = Document(out)
    notice_paragraphs = [
        p
        for p in doc.paragraphs
        if "待补充" in p.text or "内容冲突" in p.text
    ]

    assert len(notice_paragraphs) == 2
    assert all(any(run.font.highlight_color is not None for run in p.runs) for p in notice_paragraphs)
    assert all(
        p._element.pPr is not None and p._element.pPr.find(qn("w:shd")) is not None
        for p in notice_paragraphs
    )
    first_notice_runs = [run for run in notice_paragraphs[0].runs if run.text]
    assert first_notice_runs[0].text == "【待补充："
    assert first_notice_runs[0].bold is True
    assert "下一年度销量预测" in first_notice_runs[1].text
    assert first_notice_runs[1].bold is not True


def test_scalar_conflict_warns_in_validation_and_renders_readable_notice(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    chapter = _ch("c1", "产品概述", content, conflicts=["原始冲突"])
    chapter.status = "confirmed"

    validation = validate_document([chapter], None)

    assert validation["can_export"] is True
    assert validation["warnings"] == ['章节"产品概述"存在未处理冲突']

    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [chapter], template)
    doc = Document(out)

    assert any("内容冲突" in p.text and "原始冲突" in p.text for p in doc.paragraphs)


def test_malformed_issue_json_is_ignored_so_render_still_succeeds(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    chapter = _ch("c1", "产品概述", content)
    chapter.missing_information_json = "{not-json"
    chapter.conflict_json = "["

    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [chapter], template)
    doc = Document(out)

    assert any(p.text == "正文内容" for p in doc.paragraphs)
    assert not any("待补充" in p.text or "内容冲突" in p.text for p in doc.paragraphs)


def test_empty_content_falls_back_to_placeholder_text(tmp_path):
    content = {"type": "doc", "content": []}
    template = _make_template(tmp_path, ["产品概述"])
    chapter = _ch("c1", "产品概述", content)
    chapter.plain_text = ""  # no fallback text either -> must hit the placeholder
    out = render_to_docx("doc1", [chapter], template)
    doc = Document(out)
    assert any("内容待生成" in p.text for p in doc.paragraphs)


def test_no_template_builds_from_scratch_with_same_structure(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "子标题"}]},
        {"type": "paragraph", "content": [{"type": "text", "text": "正文"}]},
    ]}
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], None)
    doc = Document(out)
    assert any(p.text == "产品概述" and p.style.name == "Heading 1" for p in doc.paragraphs)
    assert any(p.text == "子标题" and p.style.name == "Heading 2" for p in doc.paragraphs)
    assert any(p.text == "正文" for p in doc.paragraphs)


def test_no_template_sets_a4_margins_and_footer_page_field(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文"}]},
    ]}

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], None)
    doc = Document(out)
    section = doc.sections[0]

    assert abs(section.page_width - Mm(210)) <= Mm(1)
    assert abs(section.page_height - Mm(297)) <= Mm(1)
    assert section.top_margin > 0
    assert section.bottom_margin > 0
    assert section.left_margin > 0
    assert section.right_margin > 0

    footer_xml = section.footer._element.xml
    assert "PAGE" in footer_xml
    assert "产品概述" in core_title(doc)
    fonts = get_export_font_config()
    normal_fonts = doc.styles["Normal"]._element.rPr.rFonts
    assert normal_fonts.get(qn("w:ascii")) == fonts.latin
    assert normal_fonts.get(qn("w:hAnsi")) == fonts.latin
    assert normal_fonts.get(qn("w:eastAsia")) == fonts.cjk
    assert normal_fonts.get(qn("w:cs")) == fonts.cjk


def test_generated_docx_styles_and_runs_use_explicit_cjk_font(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "中文正文 Latin"}]},
    ]}

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], None)
    doc = Document(out)
    fonts = get_export_font_config()
    normal_fonts = doc.styles["Normal"]._element.rPr.rFonts
    body_run = next(run for paragraph in doc.paragraphs for run in paragraph.runs if run.text == "中文正文 Latin")
    run_fonts = body_run._element.rPr.rFonts

    for r_fonts in (normal_fonts, run_fonts):
        assert r_fonts.get(qn("w:ascii")) == fonts.latin
        assert r_fonts.get(qn("w:hAnsi")) == fonts.latin
        assert r_fonts.get(qn("w:eastAsia")) == fonts.cjk
        assert r_fonts.get(qn("w:cs")) == fonts.cjk


def test_template_rendered_runs_use_explicit_cjk_font(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "模板中文正文"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    fonts = get_export_font_config()
    body_run = next(run for paragraph in doc.paragraphs for run in paragraph.runs if run.text == "模板中文正文")
    run_fonts = body_run._element.rPr.rFonts

    assert run_fonts.get(qn("w:ascii")) == fonts.latin
    assert run_fonts.get(qn("w:hAnsi")) == fonts.latin
    assert run_fonts.get(qn("w:eastAsia")) == fonts.cjk
    assert run_fonts.get(qn("w:cs")) == fonts.cjk


def test_table_rendering_adds_borders_and_header_shading(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "table", "content": [
            {"type": "tableRow", "content": [
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "序号"}]}]},
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "类别"}]}]},
            ]},
            {"type": "tableRow", "content": [
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "1"}]}]},
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "臂架系统"}]}]},
            ]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    table = doc.tables[0]

    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    assert all(borders.find(qn(f"w:{edge}")) is not None for edge in ("top", "left", "bottom", "right", "insideH", "insideV"))

    header_cell = table.rows[0].cells[0]
    tc_pr = header_cell._tc.tcPr
    assert tc_pr is not None
    assert tc_pr.find(qn("w:shd")) is not None
    assert all(run.bold for run in header_cell.paragraphs[0].runs if run.text)
    first_col_width = int(table.rows[1].cells[0]._tc.tcPr.find(qn("w:tcW")).get(qn("w:w")))
    second_col_width = int(table.rows[1].cells[1]._tc.tcPr.find(qn("w:tcW")).get(qn("w:w")))
    assert second_col_width > first_col_width


def core_title(doc: Document) -> str:
    return doc.core_properties.title or ""
