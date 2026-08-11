import json
from pathlib import Path
from unittest.mock import MagicMock
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm
from app.services.docx_renderer import render_to_docx
from app.services.export_fonts import get_export_font_config
from app.services.validator import validate_document
from app.config import SCENARIO1_TEMPLATE_PATH


REAL_TEMPLATE_PATH = SCENARIO1_TEMPLATE_PATH
ONE_PIXEL_PNG = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _ch(chapter_id, title, content_json, missing=None, conflicts=None):
    ch = MagicMock()
    ch.id = chapter_id
    ch.title = title
    ch.plain_text = "(unused fallback)"
    ch.content_json = json.dumps(content_json, ensure_ascii=False)
    ch.missing_information_json = json.dumps(missing or [], ensure_ascii=False)
    ch.conflict_json = json.dumps(conflicts or [], ensure_ascii=False)
    return ch


def _make_template(tmp_path, headings):
    """Build a minimal .docx with one Heading-1 paragraph per chapter title,
    mirroring the real target template's structure (heading, then body)."""
    doc = Document()
    for h in headings:
        doc.add_heading(h, level=1)
    path = tmp_path / "template.docx"
    doc.save(str(path))
    return str(path)


def test_image_node_is_preserved_in_exported_docx(tmp_path):
    content = {
        "type": "doc",
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "图片说明"}]},
            {"type": "image", "attrs": {"src": ONE_PIXEL_PNG, "alt": "示例图片"}},
        ],
    }
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)

    doc = Document(out)
    assert len(doc.inline_shapes) == 1


def _make_placeholder_template(tmp_path):
    doc = Document()
    doc.add_paragraph("模板封面标题")
    first_footer = doc.sections[0].footer
    first_footer.paragraphs[0].text = "模板首页页脚"
    first_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_section(WD_SECTION.NEW_PAGE)
    second_footer = doc.sections[1].footer
    second_footer.is_linked_to_previous = False
    second_footer.paragraphs[0].text = "模板正文页脚"
    second_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for title in ("第一章", "第二章"):
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"模板占位内容：{title}")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = f"模板占位表格：{title}"
    doc.add_paragraph("模板末尾说明")

    path = tmp_path / "placeholder-template.docx"
    doc.save(str(path))
    return str(path)


def _make_mixed_structure_template(tmp_path):
    doc = Document()
    doc.add_heading("第一章", level=1)
    doc.add_paragraph("模板占位内容")
    stale_table = doc.add_table(rows=1, cols=1)
    stale_table.cell(0, 0).text = "模板占位表格"
    doc.add_heading("未锚定结构标题", level=2)
    structure_table = doc.add_table(rows=1, cols=1)
    structure_table.cell(0, 0).text = "保留的结构表格"
    doc.add_heading("第二章", level=1)
    doc.add_paragraph("模板占位内容")

    path = tmp_path / "mixed-structure-template.docx"
    doc.save(str(path))
    return str(path)


def _make_template_with_appendix_tail(tmp_path):
    doc = Document()
    for title in ("第一章", "第二章"):
        doc.add_heading(title, level=1)
        doc.add_paragraph("模板占位内容")
    doc.add_heading("附录 A", level=1)
    doc.add_paragraph("附录普通尾部内容")
    appendix_table = doc.add_table(rows=1, cols=1)
    appendix_table.cell(0, 0).text = "附录结构表格"

    path = tmp_path / "appendix-tail-template.docx"
    doc.save(str(path))
    return str(path)


def _make_styled_template(tmp_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Template Body"
    for level in range(1, 4):
        doc.styles[f"Heading {level}"].font.name = f"Template Heading {level}"
    doc.add_heading("产品概述", level=1)
    path = tmp_path / "styled-template.docx"
    doc.save(str(path))
    return str(path)


def _make_literal_page_template(tmp_path):
    doc = Document()
    doc.add_heading("产品概述", level=1)
    doc.sections[0].footer.paragraphs[0].text = "PAGE"
    path = tmp_path / "literal-page-template.docx"
    doc.save(str(path))
    return str(path)


def test_heading_node_becomes_docx_heading_style(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "市场需求分析"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert any(p.text == "市场需求分析" for p in headings)


def test_bold_mark_becomes_bold_run(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "marks": [{"type": "bold"}], "text": "表6 产品卖点和优势分析表"},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
    assert any(r.text == "表6 产品卖点和优势分析表" for r in bold_runs)


def test_table_node_becomes_real_docx_table(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "table", "content": [
            {"type": "tableRow", "content": [
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "序号"}]}]},
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "类别"}]}]},
            ]},
            {"type": "tableRow", "content": [
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "1"}]}]},
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "臂架系统"}]}]},
            ]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "序号"
    assert table.rows[1].cells[1].text == "臂架系统"


def test_bullet_list_becomes_list_bullet_paragraphs(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "bulletList", "content": [
            {"type": "listItem", "content": [{"type": "paragraph", "content": [
                {"type": "text", "marks": [{"type": "bold"}], "text": "卖点"},
                {"type": "text", "text": "：主臂采用五节U型截面"},
            ]}]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    bullet_paras = [p for p in doc.paragraphs if "卖点" in p.text]
    assert len(bullet_paras) == 1
    assert "主臂采用五节U型截面" in bullet_paras[0].text


def test_ordered_list_becomes_numbered_paragraphs(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "orderedList", "content": [
            {"type": "listItem", "content": [{"type": "paragraph", "content": [
                {"type": "text", "text": "第一项"},
            ]}]},
            {"type": "listItem", "content": [{"type": "paragraph", "content": [
                {"type": "text", "text": "第二项"},
            ]}]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)

    numbered_paras = [p.text for p in doc.paragraphs if "项" in p.text]
    assert numbered_paras == ["1. 第一项", "2. 第二项"]


def test_nested_mixed_lists_render_recursively_with_indentation(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "bulletList", "content": [
            {"type": "listItem", "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "一级项目"}]},
                {"type": "orderedList", "content": [
                    {"type": "listItem", "content": [
                        {"type": "paragraph", "content": [{"type": "text", "text": "二级一"}]},
                    ]},
                    {"type": "listItem", "content": [
                        {"type": "paragraph", "content": [{"type": "text", "text": "二级二"}]},
                        {"type": "bulletList", "content": [
                            {"type": "listItem", "content": [
                                {"type": "paragraph", "content": [{"type": "text", "text": "三级"}]},
                            ]},
                        ]},
                    ]},
                ]},
            ]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)

    list_paragraphs = [p for p in doc.paragraphs if any(label in p.text for label in ("一级", "二级", "三级"))]
    assert [p.text for p in list_paragraphs] == [
        "• 一级项目",
        "1. 二级一",
        "2. 二级二",
        "▪ 三级",
    ]
    assert list_paragraphs[1].paragraph_format.left_indent > list_paragraphs[0].paragraph_format.left_indent
    assert list_paragraphs[3].paragraph_format.left_indent > list_paragraphs[1].paragraph_format.left_indent


def test_near_match_heading_uses_normalized_punctuation(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "功能性能定位生成内容"}]},
    ]}
    template = _make_template(tmp_path, ["功能/性能定位（来自技术可行性分析报告）"])

    out = render_to_docx("doc1", [_ch("c1", "功能性能定位", content)], template)
    doc = Document(out)

    assert any(p.text == "功能性能定位生成内容" for p in doc.paragraphs)


def test_chapter_anchor_requires_exact_or_delimited_suffix(tmp_path):
    def content(text):
        return {"type": "doc", "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": text}]},
        ]}
    template = _make_template(
        tmp_path,
        [
            "产品概述设计",
            "产品概述",
            "功能/性能定位（来自技术可行性分析报告）",
        ],
    )
    chapters = [
        _ch("c-overview", "产品概述", content("产品概述生成内容")),
        _ch("c-functional", "功能性能定位", content("功能性能定位生成内容")),
    ]

    out = render_to_docx("docx-anchor-boundary-regression", chapters, template)
    doc = Document(out)
    paragraph_text = [paragraph.text for paragraph in doc.paragraphs]

    assert paragraph_text.index("产品概述") < paragraph_text.index("产品概述生成内容")
    assert paragraph_text.index("产品概述设计") < paragraph_text.index("产品概述")
    assert paragraph_text.index("功能/性能定位（来自技术可行性分析报告）") < paragraph_text.index(
        "功能性能定位生成内容"
    )


def test_mixed_template_cleanup_preserves_unanchored_structure(tmp_path):
    def content(text):
        return {"type": "doc", "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": text}]},
        ]}
    chapters = [
        _ch("c1", "第一章", content("生成第一章")),
        _ch("c2", "第二章", content("生成第二章")),
    ]
    template = _make_mixed_structure_template(tmp_path)

    out = render_to_docx("doc1", chapters, template)
    doc = Document(out)
    body_text = "\n".join(p.text for p in doc.paragraphs)

    assert "生成第一章" in body_text
    assert "生成第二章" in body_text
    assert "模板占位内容" not in body_text
    assert any(p.text == "未锚定结构标题" for p in doc.paragraphs)
    assert any(cell.text == "保留的结构表格" for table in doc.tables for row in table.rows for cell in row.cells)
    assert not any(cell.text == "模板占位表格" for table in doc.tables for row in table.rows for cell in row.cells)


def test_real_template_preserves_structure_assets_and_removes_sample_body(tmp_path):
    assert REAL_TEMPLATE_PATH.exists()
    def content(text):
        return {"type": "doc", "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": text}]},
        ]}
    chapters = [
        _ch("c1", "产品概述", content("真实模板生成内容")),
        _ch("c2", "功能性能定位", content("功能性能定位生成内容")),
    ]

    out = render_to_docx("real-template-test", chapters, str(REAL_TEMPLATE_PATH))
    doc = Document(out)
    body_text = "\n".join(p.text for p in doc.paragraphs)

    assert "真实模板生成内容" in body_text
    assert "功能性能定位生成内容" in body_text
    assert any(p.text == "产品开发的必要性" for p in doc.paragraphs)
    assert any(p.text == "产品初步方案及可行性" for p in doc.paragraphs)
    assert "本产品是**规格的**产品" not in body_text
    assert "文件内容填写要求如下：" not in body_text
    assert "XXXXXXX" not in body_text
    assert len(doc.tables) == 14
    assert len(doc.sections) == 2
    assert len(doc.inline_shapes) == 3
    assert all(
        any("PAGE" in (instr.text or "").upper() for instr in section.footer._element.iter(qn("w:instrText")))
        for section in doc.sections
    )


def test_real_template_replaces_market_sample_table_with_generated_table(tmp_path):
    assert REAL_TEMPLATE_PATH.exists()
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "marks": [{"type": "bold"}], "text": "表1 产品近年销量（台数）走势及预测表"},
        ]},
        {"type": "table", "content": [
            {"type": "tableRow", "content": [
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [
                    {"type": "text", "text": "序号"},
                ]}]},
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [
                    {"type": "text", "text": "年度总销量"},
                ]}]},
            ]},
            {"type": "tableRow", "content": [
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [
                    {"type": "text", "text": "1"},
                ]}]},
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [
                    {"type": "text", "text": "100"},
                ]}]},
            ]},
        ]},
    ]}

    functional_content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "text": "功能性能定位正文"},
        ]},
    ]}
    chapters = [
        _ch("c-market", "市场需求分析", content),
        _ch("c-functional", "功能性能定位", functional_content),
    ]

    out = render_to_docx("real-template-market-table", chapters, str(REAL_TEMPLATE_PATH))
    doc = Document(out)
    table_texts = [
        "\n".join(cell.text for row in table.rows for cell in row.cells)
        for table in doc.tables
    ]

    assert any("年度总销量" in text and "100" in text for text in table_texts)
    assert not any("本年度-3年" in text or "品牌/型号1" in text for text in table_texts)
    assert sum("表1 产品近年销量（台数）走势及预测表" in p.text for p in doc.paragraphs) == 1
    assert any("客户群A" in text for text in table_texts)
    assert any("指标项次" in text for text in table_texts)
    assert any("实际解决时间" in text for text in table_texts)
    assert len(doc.tables) == 14
    assert len(doc.inline_shapes) == 3
    assert len(doc.sections) == 2
    assert all(
        any("PAGE" in (instr.text or "").upper() for instr in section.footer._element.iter(qn("w:instrText")))
        for section in doc.sections
    )


def test_literal_page_text_does_not_count_as_page_field(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文"}]},
    ]}
    template = _make_literal_page_template(tmp_path)

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    footer = doc.sections[0].footer
    instructions = [instr.text or "" for instr in footer._element.iter(qn("w:instrText"))]

    assert "PAGE" in footer.paragraphs[0].text
    assert any("PAGE" in instruction.upper() for instruction in instructions)
    assert any(fld.get(qn("w:fldCharType")) == "begin" for fld in footer._element.iter(qn("w:fldChar")))


def test_template_cleanup_preserves_structure_and_footer_page_fields(tmp_path):
    def content(text):
        return {"type": "doc", "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": text}]},
        ]}
    chapters = [
        _ch("c1", "第一章", content("生成第一章")),
        _ch("c2", "第二章", content("生成第二章")),
    ]
    template = _make_placeholder_template(tmp_path)

    out = render_to_docx("doc1", chapters, template)
    doc = Document(out)
    body_text = "\n".join(p.text for p in doc.paragraphs)

    assert "模板封面标题" in body_text
    assert "生成第一章" in body_text
    assert "生成第二章" in body_text
    assert "模板占位内容" not in body_text
    assert "模板末尾说明" not in body_text
    assert all(
        any(p.text == title and p.style.name == "Heading 1" for p in doc.paragraphs)
        for title in ("第一章", "第二章")
    )
    assert len(doc.tables) == 0

    first_footer = doc.sections[0].footer
    second_footer = doc.sections[1].footer
    first_footer_text = "\n".join(p.text for p in first_footer.paragraphs)
    second_footer_text = "\n".join(p.text for p in second_footer.paragraphs)
    assert "模板首页页脚" in first_footer_text
    assert "模板正文页脚" in second_footer_text
    assert first_footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert second_footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert "PAGE" in first_footer._element.xml
    assert "PAGE" in second_footer._element.xml


def test_missing_information_stays_out_of_exported_docx(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx(
        "doc1", [_ch("c1", "产品概述", content, missing=["下一年度销量预测"])], template
    )
    doc = Document(out)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "正文内容" in body_text
    assert "待补充" not in body_text
    assert "下一年度销量预测" not in body_text


def test_conflict_notes_are_rendered_as_highlighted_notices(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx(
        "doc1",
        [_ch(
            "c1",
            "产品概述",
            content,
            conflicts=[{"description": "销售目标与产能规划不一致"}],
        )],
        template,
    )

    doc = Document(out)
    notice_paragraphs = [
        p
        for p in doc.paragraphs
        if "内容冲突" in p.text
    ]

    assert len(notice_paragraphs) == 1
    assert all(any(run.font.highlight_color is not None for run in p.runs) for p in notice_paragraphs)
    assert all(
        p._element.pPr is not None and p._element.pPr.find(qn("w:shd")) is not None
        for p in notice_paragraphs
    )
    conflict_runs = [run for run in notice_paragraphs[0].runs if run.text]
    assert conflict_runs[0].text == "【内容冲突："
    assert conflict_runs[0].bold is True
    assert "销售目标与产能规划不一致" in conflict_runs[1].text
    assert conflict_runs[1].bold is not True


def test_scalar_conflict_warns_in_validation_and_renders_readable_notice(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    chapter = _ch("c1", "产品概述", content, conflicts=["原始冲突"])
    chapter.status = "confirmed"

    validation = validate_document([chapter], None)

    assert validation["can_export"] is True
    assert validation["warnings"] == ['章节"产品概述"存在未处理冲突']

    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [chapter], template)
    doc = Document(out)

    assert any("内容冲突" in p.text and "原始冲突" in p.text for p in doc.paragraphs)


def test_malformed_issue_json_is_ignored_so_render_still_succeeds(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文内容"}]},
    ]}
    chapter = _ch("c1", "产品概述", content)
    chapter.missing_information_json = "{not-json"
    chapter.conflict_json = "["

    template = _make_template(tmp_path, ["产品概述"])
    out = render_to_docx("doc1", [chapter], template)
    doc = Document(out)

    assert any(p.text == "正文内容" for p in doc.paragraphs)
    assert not any("待补充" in p.text or "内容冲突" in p.text for p in doc.paragraphs)


def test_empty_content_falls_back_to_placeholder_text(tmp_path):
    content = {"type": "doc", "content": []}
    template = _make_template(tmp_path, ["产品概述"])
    chapter = _ch("c1", "产品概述", content)
    chapter.plain_text = ""  # no fallback text either -> must hit the placeholder
    out = render_to_docx("doc1", [chapter], template)
    doc = Document(out)
    assert any("内容待生成" in p.text for p in doc.paragraphs)


def test_non_dict_json_roots_fall_back_to_plain_text(tmp_path):
    template = _make_template(tmp_path, ["产品概述"])

    for index, root in enumerate((None, "not-a-document", [], 7)):
        chapter = _ch(f"c{index}", "产品概述", root)
        chapter.plain_text = f"根节点回退内容 {index}"

        out = render_to_docx(f"docx-root-boundary-{index}", [chapter], template)
        doc = Document(out)

        assert any(p.text == chapter.plain_text for p in doc.paragraphs)


def test_invalid_content_shapes_fall_back_to_plain_text(tmp_path):
    template = _make_template(tmp_path, ["产品概述"])

    for index, content in enumerate((None, "not-a-node-list", [None, "not-a-node", 7])):
        chapter = _ch(
            f"c{index}",
            "产品概述",
            {"type": "doc", "content": content},
        )
        chapter.plain_text = f"内容回退文本 {index}"

        out = render_to_docx(f"docx-content-boundary-{index}", [chapter], template)
        doc = Document(out)

        assert any(p.text == chapter.plain_text for p in doc.paragraphs)


def test_invalid_content_nodes_are_filtered_without_losing_valid_nodes(tmp_path):
    content = {
        "type": "doc",
        "content": [
            None,
            {"type": "paragraph", "content": [{"type": "text", "text": "有效正文"}]},
            "not-a-node",
            7,
        ],
    }
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("docx-node-filter-boundary", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)

    assert any(p.text == "有效正文" for p in doc.paragraphs)
    assert not any(p.text == "(unused fallback)" for p in doc.paragraphs)


def test_all_invalid_content_nodes_use_empty_placeholder(tmp_path):
    content = {"type": "doc", "content": [None, "not-a-node", 7]}
    template = _make_template(tmp_path, ["产品概述"])
    chapter = _ch("c1", "产品概述", content)
    chapter.plain_text = ""

    out = render_to_docx("docx-invalid-node-placeholder", [chapter], template)
    doc = Document(out)

    assert any(p.text == "（内容待生成）" for p in doc.paragraphs)


def test_malformed_nested_nodes_fall_back_to_plain_text(tmp_path):
    content = {
        "type": "doc",
        "content": [
            {"type": "paragraph", "attrs": None, "content": None},
            {
                "type": "bulletList",
                "attrs": "wrong",
                "content": [None, "not-a-list-item", {"type": "listItem", "content": None}],
            },
            {
                "type": "table",
                "attrs": [],
                "content": [
                    None,
                    "not-a-row",
                    {"type": "tableRow", "content": None},
                    {
                        "type": "tableRow",
                        "content": [None, "not-a-cell", {"type": "tableCell", "content": None}],
                    },
                ],
            },
        ],
    }
    chapter = _ch("c1", "产品概述", content)
    chapter.plain_text = "嵌套形状回退正文"
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("docx-nested-malformed-fallback", [chapter], template)
    doc = Document(out)

    assert any(p.text == chapter.plain_text for p in doc.paragraphs)
    assert len(doc.tables) == 0


def test_malformed_nested_nodes_are_filtered_while_valid_content_renders(tmp_path):
    content = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": "wrong"},
                "content": [None, "not-a-run", {"type": "text", "text": "有效标题"}],
            },
            {
                "type": "paragraph",
                "attrs": None,
                "content": [
                    None,
                    "not-a-run",
                    {"type": "text", "text": "有效正文", "marks": None},
                    {"type": "text", "text": 7},
                ],
            },
            {
                "type": "bulletList",
                "attrs": "wrong",
                "content": [
                    None,
                    "not-a-list-item",
                    {
                        "type": "listItem",
                        "attrs": None,
                        "content": [
                            None,
                            {
                                "type": "paragraph",
                                "attrs": [],
                                "content": [
                                    None,
                                    "not-a-run",
                                    {
                                        "type": "text",
                                        "text": "有效列表",
                                        "marks": [None, "wrong", {"type": "bold"}],
                                    },
                                ],
                            },
                        ],
                    },
                ],
            },
            {
                "type": "table",
                "attrs": None,
                "content": [
                    None,
                    "not-a-row",
                    {
                        "type": "tableRow",
                        "content": [
                            None,
                            "not-a-cell",
                            {
                                "type": "tableHeader",
                                "attrs": "wrong",
                                "content": [
                                    None,
                                    {
                                        "type": "paragraph",
                                        "content": [None, {"type": "text", "text": "表头"}],
                                    },
                                ],
                            },
                            {
                                "type": "tableCell",
                                "attrs": [],
                                "content": [
                                    None,
                                    "not-a-block",
                                    {
                                        "type": "paragraph",
                                        "content": [
                                            None,
                                            {"type": "text", "text": "表格值", "marks": None},
                                        ],
                                    },
                                ],
                            },
                        ],
                    },
                ],
            },
        ],
    }
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("docx-nested-malformed-filter", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)

    assert any(p.text == "有效标题" and p.style.name == "Heading 1" for p in doc.paragraphs)
    assert any(p.text == "有效正文" for p in doc.paragraphs)
    assert any(p.text == "• 有效列表" for p in doc.paragraphs)
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "表头"
    assert doc.tables[0].rows[0].cells[1].text == "表格值"


def test_no_template_builds_from_scratch_with_same_structure(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "子标题"}]},
        {"type": "paragraph", "content": [{"type": "text", "text": "正文"}]},
    ]}
    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], None)
    doc = Document(out)
    assert any(p.text == "产品概述" and p.style.name == "Heading 1" for p in doc.paragraphs)
    assert any(p.text == "子标题" and p.style.name == "Heading 2" for p in doc.paragraphs)
    assert any(p.text == "正文" for p in doc.paragraphs)


def test_no_template_sets_a4_margins_and_footer_page_field(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "正文"}]},
    ]}

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], None)
    doc = Document(out)
    section = doc.sections[0]

    assert abs(section.page_width - Mm(210)) <= Mm(1)
    assert abs(section.page_height - Mm(297)) <= Mm(1)
    assert section.top_margin > 0
    assert section.bottom_margin > 0
    assert section.left_margin > 0
    assert section.right_margin > 0

    footer_xml = section.footer._element.xml
    assert "PAGE" in footer_xml
    assert "产品概述" in core_title(doc)
    fonts = get_export_font_config()
    normal_fonts = doc.styles["Normal"]._element.rPr.rFonts
    assert normal_fonts.get(qn("w:ascii")) == fonts.latin
    assert normal_fonts.get(qn("w:hAnsi")) == fonts.latin
    assert normal_fonts.get(qn("w:eastAsia")) == fonts.cjk
    assert normal_fonts.get(qn("w:cs")) == fonts.cjk


def test_generated_docx_styles_and_runs_use_explicit_cjk_font(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "中文正文 Latin"}]},
    ]}

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], None)
    doc = Document(out)
    fonts = get_export_font_config()
    normal_fonts = doc.styles["Normal"]._element.rPr.rFonts
    body_run = next(run for paragraph in doc.paragraphs for run in paragraph.runs if run.text == "中文正文 Latin")
    run_fonts = body_run._element.rPr.rFonts

    for r_fonts in (normal_fonts, run_fonts):
        assert r_fonts.get(qn("w:ascii")) == fonts.latin
        assert r_fonts.get(qn("w:hAnsi")) == fonts.latin
        assert r_fonts.get(qn("w:eastAsia")) == fonts.cjk
        assert r_fonts.get(qn("w:cs")) == fonts.cjk


def test_template_rendered_runs_use_explicit_cjk_font(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "模板中文正文"}]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    fonts = get_export_font_config()
    body_run = next(run for paragraph in doc.paragraphs for run in paragraph.runs if run.text == "模板中文正文")
    run_fonts = body_run._element.rPr.rFonts

    assert run_fonts.get(qn("w:ascii")) == fonts.latin
    assert run_fonts.get(qn("w:hAnsi")) == fonts.latin
    assert run_fonts.get(qn("w:eastAsia")) == fonts.cjk
    assert run_fonts.get(qn("w:cs")) == fonts.cjk


def test_table_rendering_adds_borders_and_header_shading(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "table", "content": [
            {"type": "tableRow", "content": [
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "序号"}]}]},
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "类别"}]}]},
            ]},
            {"type": "tableRow", "content": [
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "1"}]}]},
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "臂架系统"}]}]},
            ]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("doc1", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    table = doc.tables[0]

    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    assert all(borders.find(qn(f"w:{edge}")) is not None for edge in ("top", "left", "bottom", "right", "insideH", "insideV"))

    header_cell = table.rows[0].cells[0]
    tc_pr = header_cell._tc.tcPr
    assert tc_pr is not None
    assert tc_pr.find(qn("w:shd")) is not None
    assert all(run.bold for run in header_cell.paragraphs[0].runs if run.text)
    first_col_width = int(table.rows[1].cells[0]._tc.tcPr.find(qn("w:tcW")).get(qn("w:w")))
    second_col_width = int(table.rows[1].cells[1]._tc.tcPr.find(qn("w:tcW")).get(qn("w:w")))
    assert second_col_width > first_col_width


def test_generated_table_width_xml_uses_word_dxa_twips(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "table", "content": [
            {"type": "tableRow", "content": [
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "编号"}]}]},
                {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "较长的说明列"}]}]},
            ]},
            {"type": "tableRow", "content": [
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "1"}]}]},
                {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "正文"}]}]},
            ]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("docx-width-regression", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    row_widths = [
        table_cell._tc.tcPr.find(qn("w:tcW"))
        for table_cell in doc.tables[0].rows[1].cells
    ]

    assert all(width.get(qn("w:type")) == "dxa" for width in row_widths)
    assert sum(int(width.get(qn("w:w"))) for width in row_widths) == Cm(16.5).twips
    assert all(0 < int(width.get(qn("w:w"))) < 10000 for width in row_widths)


def test_wide_generated_table_keeps_all_cell_widths_positive(tmp_path):
    def cell(cell_type, text):
        return {
            "type": cell_type,
            "content": [{"type": "paragraph", "content": [{"type": "text", "text": text}]}],
        }

    content = {"type": "doc", "content": [
        {"type": "table", "content": [
            {"type": "tableRow", "content": [
                cell("tableHeader", f"列{i + 1}") for i in range(7)
            ]},
            {"type": "tableRow", "content": [
                cell("tableCell", f"值{i + 1}") for i in range(7)
            ]},
        ]},
    ]}
    template = _make_template(tmp_path, ["产品概述"])

    out = render_to_docx("docx-wide-table-regression", [_ch("c1", "产品概述", content)], template)
    doc = Document(out)
    widths = [
        cell._tc.tcPr.find(qn("w:tcW"))
        for row in doc.tables[0].rows
        for cell in row.cells
    ]

    assert all(width.get(qn("w:type")) == "dxa" for width in widths)
    assert all(0 < int(width.get(qn("w:w"))) < 10000 for width in widths)
    assert sum(int(width.get(qn("w:w"))) for width in widths[:7]) == Cm(16.5).twips


def test_template_cleanup_retains_ordinary_appendix_tail_content(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "生成内容"}]},
    ]}
    template = _make_template_with_appendix_tail(tmp_path)

    out = render_to_docx(
        "docx-appendix-tail-regression",
        [_ch("c1", "第一章", content), _ch("c2", "第二章", content)],
        template,
    )
    doc = Document(out)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "模板占位内容" not in body_text
    assert "附录 A" in body_text
    assert "附录普通尾部内容" in body_text
    assert any(
        cell.text == "附录结构表格"
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )


def test_template_rendering_preserves_normal_and_heading_styles(tmp_path):
    content = {"type": "doc", "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "生成内容"}]},
    ]}
    template = _make_styled_template(tmp_path)

    out = render_to_docx(
        "docx-template-style-regression",
        [_ch("c1", "产品概述", content)],
        template,
    )
    doc = Document(out)

    assert doc.styles["Normal"].font.name == "Template Body"
    for level in range(1, 4):
        assert doc.styles[f"Heading {level}"].font.name == f"Template Heading {level}"
    generated_run = next(
        run
        for paragraph in doc.paragraphs
        for run in paragraph.runs
        if run.text == "生成内容"
    )
    assert generated_run._element.rPr.rFonts.get(qn("w:eastAsia")) == get_export_font_config().cjk


def core_title(doc: Document) -> str:
    return doc.core_properties.title or ""
