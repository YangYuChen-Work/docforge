import json
import zipfile
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from openpyxl import load_workbook
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.config import settings
from app.db.models import Annotation, DocumentChapter, DocumentTemplate, GeneratedDocument
from app.db.session import Base
from app.domain import exports as domain_exports
from app.services.docx_renderer import render_to_docx
from app.services.xlsx_exporter import export_tables_to_xlsx


def _chapter(title="产品概述"):
    return SimpleNamespace(
        id="chapter-1",
        title=title,
        plain_text="关键内容",
        content_json=json.dumps(
            {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": "这是关键内容。"}],
                    },
                    {
                        "type": "table",
                        "content": [
                            {
                                "type": "tableRow",
                                "content": [
                                    {
                                        "type": "tableHeader",
                                        "content": [
                                            {
                                                "type": "paragraph",
                                                "content": [{"type": "text", "text": "指标"}],
                                            }
                                        ],
                                    }
                                ],
                            },
                            {
                                "type": "tableRow",
                                "content": [
                                    {
                                        "type": "tableCell",
                                        "content": [
                                            {
                                                "type": "paragraph",
                                                "content": [{"type": "text", "text": "目标值"}],
                                            }
                                        ],
                                    }
                                ],
                            },
                        ],
                    },
                ],
            },
            ensure_ascii=False,
        ),
        missing_information_json="[]",
        conflict_json="[]",
    )


def _annotation(target_text="关键内容", content="请核对这一项"):
    return SimpleNamespace(
        id="annotation-1",
        chapter_id="chapter-1",
        chapter_title="产品概述",
        target_text=target_text,
        content=content,
        status="pending",
        created_by="审阅人",
        label="审阅",
        locator="正文第1段",
    )


def _template(tmp_path):
    path = Path(tmp_path) / "template.docx"
    document = Document()
    document.add_heading("产品概述", level=1)
    document.save(path)
    return str(path)


def test_docx_export_uses_native_comments_only_when_enabled(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "storage_root", str(tmp_path))
    chapter = _chapter()
    annotation = _annotation()

    without_comments = render_to_docx(
        "doc-without-comments",
        [chapter],
        _template(tmp_path),
        annotations=[annotation],
        include_comments=False,
    )
    with zipfile.ZipFile(without_comments) as archive:
        assert "word/comments.xml" not in archive.namelist()

    with_comments = render_to_docx(
        "doc-with-comments",
        [chapter],
        _template(tmp_path),
        annotations=[annotation],
        include_comments=True,
    )
    with zipfile.ZipFile(with_comments) as archive:
        comments_xml = archive.read("word/comments.xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert comments_xml.count("<w:comment ") == 1
    assert "请核对这一项" in comments_xml
    assert "w:commentRangeStart" in document_xml
    assert "w:commentRangeEnd" in document_xml
    assert "w:commentReference" in document_xml
    document = Document(with_comments)
    assert "关键内容" in document.paragraphs[-1].text or any(
        "关键内容" in paragraph.text for paragraph in document.paragraphs
    )


def test_pdf_style_export_places_visible_comment_note_near_chapter(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "storage_root", str(tmp_path))
    out = render_to_docx(
        "pdf-comments",
        [_chapter()],
        _template(tmp_path),
        annotations=[_annotation()],
        include_comments=True,
        comment_mode="visible",
    )

    document = Document(out)
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "批注" in body
    assert "请核对这一项" in body
    with zipfile.ZipFile(out) as archive:
        assert "word/comments.xml" not in archive.namelist()


def test_xlsx_export_attaches_comments_to_matching_cells_and_lists_all_comments(
    tmp_path, monkeypatch
):
    monkeypatch.setattr(settings, "storage_root", str(tmp_path))
    out = export_tables_to_xlsx(
        "xlsx-comments",
        [_chapter()],
        annotations=[_annotation("目标值", "请确认目标值来源")],
    )

    workbook = load_workbook(out)
    assert "批注" in workbook.sheetnames
    table_sheet = workbook["产品概述"]
    assert table_sheet["A4"].comment is not None
    assert table_sheet["A4"].comment.text == "请确认目标值来源"

    comments_sheet = workbook["批注"]
    assert comments_sheet["A1"].value == "批注清单"
    assert comments_sheet["A4"].value == "目标值"
    assert comments_sheet["B4"].value == "请确认目标值来源"
    assert comments_sheet["F4"].value == "产品概述!A4"


def test_xlsx_export_keeps_unmatched_comment_in_comment_list(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "storage_root", str(tmp_path))
    out = export_tables_to_xlsx(
        "xlsx-unmatched-comments",
        [_chapter()],
        annotations=[_annotation("正文中不存在的文字", "请补充依据")],
    )

    workbook = load_workbook(out)
    comments_sheet = workbook["批注"]
    assert comments_sheet["F4"].value == "未在表格中定位"


def test_domain_persists_include_comments_and_passes_annotations_to_renderer(tmp_path, monkeypatch):
    engine = create_engine(
        f"sqlite:///{tmp_path / 'export-comments.db'}",
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(bind=engine)
    session_factory = sessionmaker(bind=engine, autocommit=False, autoflush=False)
    db = session_factory()
    db.add_all(
        [
            DocumentTemplate(
                id="TPL-C",
                name="模板",
                phase="phase",
                category="category",
                source_path=None,
            ),
            GeneratedDocument(
                id="DOC-C",
                project_id="P001",
                generation_task_id="TASK-C",
                template_id="TPL-C",
                title="批注文档",
                status="confirmed",
            ),
            DocumentChapter(
                id="CH-C",
                document_id="DOC-C",
                title="产品概述",
                order_index=1,
                status="confirmed",
                content_json=json.dumps({"type": "doc", "content": []}),
            ),
            Annotation(
                id="ANN-C",
                chapter_id="CH-C",
                type="review",
                label="审阅",
                target_text="关键内容",
                content="请核对",
                status="pending",
                created_by="审阅人",
            ),
        ]
    )
    db.commit()

    monkeypatch.setattr(
        domain_exports,
        "validate_document",
        lambda chapters, template_path: {
            "passed": True,
            "warnings": [],
            "errors": [],
            "can_export": True,
            "has_missing_info": False,
        },
    )
    captured = {}

    def fake_render(doc_id, chapters, template_path, **kwargs):
        captured.update(kwargs)
        output = tmp_path / "comments.docx"
        output.write_bytes(b"fake")
        return str(output)

    monkeypatch.setattr(domain_exports, "render_to_docx", fake_render)
    export = domain_exports.create_export(db, "DOC-C", "docx", include_comments=True)

    assert export.status == "completed"
    assert export.include_comments is True
    assert captured["include_comments"] is True
    assert captured["comment_mode"] == "native"
    assert captured["annotations"][0]["content"] == "请核对"
    db.close()
