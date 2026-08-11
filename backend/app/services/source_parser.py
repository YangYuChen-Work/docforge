import json
from pathlib import Path
from docx import Document as DocxDocument
from openpyxl import load_workbook


def parse_docx(file_path: str) -> tuple[list[dict], int]:
    """Parse DOCX file. Returns (content_items, image_count)."""
    doc = DocxDocument(file_path)
    items = []
    order = 0
    heading_stack: list[tuple[int, str]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name

        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            heading_stack = [(l, t) for l, t in heading_stack if l < level]
            heading_stack.append((level, text))
            heading_path = " / ".join(t for _, t in heading_stack)
            items.append({
                "content_type": "heading",
                "heading_level": level,
                "heading_path": heading_path,
                "content_text": text,
                "structured_value": None,
                "locator": heading_path,
                "order_index": order,
            })
        else:
            items.append({
                "content_type": "paragraph",
                "heading_level": None,
                "heading_path": " / ".join(t for _, t in heading_stack),
                "content_text": text,
                "structured_value": None,
                "locator": None,
                "order_index": order,
            })
        order += 1

    for t_idx, table in enumerate(doc.tables):
        headers = []
        rows = []
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if r_idx == 0:
                headers = cells
            else:
                rows.append(cells)
        items.append({
            "content_type": "table",
            "heading_level": None,
            "heading_path": " / ".join(t for _, t in heading_stack),
            "content_text": " | ".join(headers),
            "structured_value": json.dumps(
                {"headers": headers, "rows": rows[:50]}, ensure_ascii=False
            ),
            "locator": f"表{t_idx + 1}",
            "order_index": order,
        })
        order += 1

    # Count images
    image_count = sum(
        1 for rel in doc.part.rels.values() if "image" in rel.reltype
    )

    return items, image_count


def parse_xlsx(file_path: str) -> tuple[list[dict], int]:
    """Parse XLSX file. Returns (content_items, 0)."""
    wb = load_workbook(file_path, read_only=True, data_only=True)
    items = []
    order = 0
    for sheet in wb.worksheets:
        headers: list[str] = []
        rows: list[list[str]] = []
        for r_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            if all(c == "" for c in cells):
                continue
            if r_idx == 0:
                headers = cells
            else:
                rows.append(cells)
        items.append({
            "content_type": "table",
            "heading_level": None,
            "heading_path": sheet.title,
            "content_text": " | ".join(headers),
            "structured_value": json.dumps(
                {"sheet": sheet.title, "headers": headers, "rows": rows[:100]},
                ensure_ascii=False,
            ),
            "locator": f"工作表:{sheet.title}",
            "order_index": order,
        })
        order += 1
    wb.close()
    return items, 0
